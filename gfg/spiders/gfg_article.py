import scrapy
from docx import Document

doc = Document()

class GfgArticleSpider(scrapy.Spider):
    name = 'gfg_article'
    allowed_domains = ['www.geeksforgeeks.org']
    article_url = 'https://www.geeksforgeeks.org/data-structures/?ref=shm'
    start_urls = [article_url]

    def parse(self, response):
        
        title = response.css("h1.entry-title::text").get()
        page_content = response.css('div.page_content *::text').getall()

        doc.add_heading(title)
        
        for c in page_content:
            doc.add_paragraph(c)
        
        doc.save('output.docx')


